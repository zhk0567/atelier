#!/usr/bin/env python3
"""Apply paragraph rewrites to existing homework docx in-place (preserve formatting)."""

from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
SCRIPTS = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPTS))

from rewrite_map import DUPLICATE_DELETE_MATCH, REWRITES  # noqa: E402

HOMEWORK = ROOT / "作业"
OUT = HOMEWORK / "B23070426_张浩坤_基于FastAPI的个人站点在阿里云ECS上的部署与实现.docx"
BACKUP = HOMEWORK / "backup"

SKIP_PREFIXES = (
    "图",
    "表",
    "指 导 教 师 评 语",
    "参考文献",
    "目录",
    "云计算技术课程大作业",
    "题目：",
    "学    号",
    "专    业",
    "学生姓名",
    "任课教师",
    "完成度评价",
    "成绩：",
    "指导教师",
    "洛阳理工学院",
    "B23070426",
    "第1章 绪论\t",
    "      第",
    "      1.",
    "      2.",
    "      3.",
    "      4.",
    "      5.",
    "      6.",
    "      7.",
)

SKIP_EXACT = {
    "张浩坤",
    "数据科学与大数据技术",
    "龚蕾",
    "_______________",
}


def should_skip(text: str) -> bool:
    t = text.strip()
    if not t or len(t) < 4:
        return True
    if t in SKIP_EXACT:
        return True
    for p in SKIP_PREFIXES:
        if t.startswith(p):
            return True
    if re.match(r"^[\[\d]", t):
        return True
    if re.match(r"^T\d\b", t):
        return True
    return False


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def backup_docx(src: Path) -> Path:
    BACKUP.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest = BACKUP / f"{src.stem}_pre_apply_{ts}.docx"
    shutil.copy2(src, dest)
    return dest


def remove_duplicate_boilerplate(doc) -> int:
    seen = False
    removed = 0
    for para in list(doc.paragraphs):
        t = para.text.strip()
        if DUPLICATE_DELETE_MATCH not in t:
            continue
        if not seen:
            seen = True
            continue
        delete_paragraph(para)
        removed += 1
    return removed


def apply_rewrites(doc) -> tuple[int, int]:
    replaced = 0
    skipped_once: set[str] = set()
    for para in doc.paragraphs:
        old = para.text.strip()
        if should_skip(old):
            continue
        current = old
        changed = False
        for rule in REWRITES:
            match = rule["match"]
            if match not in current:
                continue
            if rule.get("once") and match in skipped_once:
                continue
            mode = rule.get("mode", "substring")
            if mode == "full":
                current = rule["new"]
            else:
                current = current.replace(match, rule["new"], 1)
            changed = True
            replaced += 1
            if rule.get("once"):
                skipped_once.add(match)
        if changed:
            set_paragraph_text(para, current)
    return replaced, len(skipped_once)


def count_figures(doc) -> int:
    return sum(1 for p in doc.paragraphs if p.text.strip().startswith("图"))


def main() -> None:
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx")
        raise SystemExit(1)

    if not OUT.is_file():
        print("missing", OUT)
        raise SystemExit(1)

    pre = backup_docx(OUT)
    print("backup:", pre)

    doc = Document(str(OUT))
    para_before = len(doc.paragraphs)
    fig_before = count_figures(doc)

    removed = remove_duplicate_boilerplate(doc)
    replaced, _ = apply_rewrites(doc)

    para_after = len(doc.paragraphs)
    fig_after = count_figures(doc)

    try:
        doc.save(str(OUT))
        print("saved:", OUT)
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_降重已更新.docx")
        doc.save(str(alt))
        print("原文件被 Word 占用，已另存为:", alt)
        print("请关闭 Word 后覆盖原 docx。")

    print(f"paragraphs: {para_before} -> {para_after} (removed {removed} duplicates)")
    print(f"figure captions: {fig_before} -> {fig_after}")
    print(f"rewritten paragraphs: {replaced}")
    print(f"rewrite rules loaded: {len(REWRITES)}")


if __name__ == "__main__":
    main()
