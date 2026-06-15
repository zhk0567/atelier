#!/usr/bin/env python3
"""Build expanded homework docx from report_content + images extracted from existing docx."""

from __future__ import annotations

import shutil
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
SCRIPTS = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPTS))

from report_content import COVER, FIGURE_TITLES, SHOT_FILES, get_blocks  # noqa: E402

HOMEWORK = ROOT / "作业"
OUT = HOMEWORK / "B23070426_张浩坤_基于FastAPI的个人站点在阿里云ECS上的部署与实现.docx"
SHOTS = HOMEWORK / "_shots_tmp"


def extract_images(src: Path) -> None:
    SHOTS.mkdir(parents=True, exist_ok=True)
    if not src.is_file():
        return
    with zipfile.ZipFile(src) as z:
        media = sorted(n for n in z.namelist() if n.startswith("word/media/") and not n.endswith("/"))
        for i, name in enumerate(media[: len(SHOT_FILES)]):
            (SHOTS / SHOT_FILES[i]).write_bytes(z.read(name))


def main() -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
    except ImportError:
        print("pip install python-docx")
        raise SystemExit(1)

    old = OUT if OUT.is_file() else next(HOMEWORK.glob("B23070426*.docx"), None)
    if old:
        extract_images(old)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    def cen(text: str, size: int = 14, bold: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.name = "宋体"
        r.font.size = Pt(size)
        r.bold = bold
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def heading(level: int, text: str) -> None:
        if level == 1:
            doc.add_page_break()
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def body(text: str) -> None:
        if text.strip():
            doc.add_paragraph(text.strip())

    def table(caption: str, headers: list[str], rows: list[list[str]]) -> None:
        if caption:
            p = doc.add_paragraph(caption)
            if p.runs:
                p.runs[0].bold = True
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                t.rows[ri + 1].cells[ci].text = val

    def code(text: str) -> None:
        p = doc.add_paragraph(text)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(9)

    def figure(ch: int, idx: int) -> None:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(3)
        title = FIGURE_TITLES[idx] if idx < len(FIGURE_TITLES) else ""
        r = cap.add_run(f"图{ch}-{idx + 1}  {title}")
        r.font.name = "楷体"
        r.font.size = Pt(10.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
        path = SHOTS / SHOT_FILES[idx]
        if path.is_file():
            doc.add_picture(str(path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # cover
    cen("云计算技术课程大作业", 18, True)
    cen(f"题目：{COVER['title']}", 15)
    cen(f"学    号：{COVER['student_id']}", 14)
    cen(f"专    业：{COVER['major']}", 14)
    cen(f"学生姓名：{COVER['name']}", 14)
    cen("任课教师：", 14)
    cen(f"{COVER['school']}    2025年9月", 14)

    for block in get_blocks():
        kind = block[0]
        if kind == "h1":
            heading(1, block[1])
        elif kind == "h2":
            heading(2, block[1])
        elif kind == "h3":
            heading(3, block[1])
        elif kind == "body":
            body(block[1])
        elif kind == "table":
            _, caption, headers, rows = block
            table(caption, headers, rows)
        elif kind == "code":
            code(block[1])
        elif kind == "figures":
            _, ch, indices = block
            for idx in indices:
                figure(ch, idx)

    doc.add_page_break()
    cen("指 导 教 师 评 语", 16, True)
    body("完成度评价：")
    body("成绩：")
    cen("指导教师：_______________    2026年6月20日", 14)

    HOMEWORK.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)

    if SHOTS.is_dir():
        shutil.rmtree(SHOTS, ignore_errors=True)

    text_len = sum(len(p.text) for p in doc.paragraphs)
    print(f"已生成: {OUT}")
    print(f"正文字符约: {text_len}")
    print("请在 Word 中确认页数 ≥35；图5-1～5-10 可替换为阿里云真实截图。")


if __name__ == "__main__":
    main()
