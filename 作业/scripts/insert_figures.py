#!/usr/bin/env python3
"""Insert figure captions and images into existing homework docx without rebuilding."""

from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
HOMEWORK = ROOT / "作业"
OUT = HOMEWORK / "B23070426_张浩坤_基于FastAPI的个人站点在阿里云ECS上的部署与实现.docx"
SHOTS = HOMEWORK / "screenshots"
BACKUP = HOMEWORK / "backup"

# anchor heading substring -> list of figures to insert at end of that section
FIGURE_SLOTS: list[dict] = [
    {
        "anchor": "1.2 云计算服务模型与选型",
        "figures": [
            (1, 1, "云计算三层服务模型对比示意", SHOTS / "ch1" / "01-服务模型.png"),
        ],
    },
    {
        "anchor": "1.3",
        "anchor_alt": "Web 部署的关键环节",
        "figures": [
            (1, 2, "个人站点建设动机与课程对应关系", SHOTS / "ch1" / "02-建设动机.png"),
        ],
    },
    {
        "anchor": "2.3 总体架构",
        "figures": [
            (2, 1, "atelier 总体部署架构", SHOTS / "ch2" / "01-总体架构.png"),
            (2, 2, "单次 HTTPS 请求生命周期", SHOTS / "ch2" / "02-请求生命周期.png"),
        ],
    },
    {
        "anchor": "3.1 全站布局与导航",
        "figures": [(3, 1, "全站顶栏与侧栏导航", SHOTS / "ch3" / "01-全站导航.png")],
    },
    {
        "anchor": "3.2 首页",
        "figures": [(3, 2, "首页置顶项目与博客入口", SHOTS / "ch3" / "02-首页.png")],
    },
    {
        "anchor": "3.3 项目展示模块",
        "figures": [(3, 3, "项目列表或详情页", SHOTS / "ch3" / "03-项目详情.png")],
    },
    {
        "anchor": "3.4 博客模块",
        "figures": [(3, 4, "Framework 博客系列页", SHOTS / "ch3" / "04-博客系列.png")],
    },
    {
        "anchor": "3.5 Wiki 模块",
        "figures": [(3, 5, "Wiki 分页与侧栏导航", SHOTS / "ch3" / "05-Wiki分页.png")],
    },
    {
        "anchor": "3.6 数据浏览 browse",
        "figures": [(3, 6, "browse 数据表格页", SHOTS / "ch3" / "06-browse.png")],
    },
    {
        "anchor": "3.7 旅行相册",
        "figures": [(3, 7, "旅行相册列表或详情", SHOTS / "ch3" / "07-旅行相册.png")],
    },
    {
        "anchor": "3.8 壁纸、Live2D 与 API",
        "figures": [(3, 8, "壁纸切换或 Live2D 展示", SHOTS / "ch3" / "08-壁纸Live2D.png")],
    },
    {
        "anchor": "4.1 应用入口与路由",
        "figures": [(4, 1, "FastAPI 路由模块划分", SHOTS / "ch4" / "01-模块划分.png")],
    },
    {
        "anchor": "4.2 数据流与渲染",
        "figures": [(4, 2, "数据流：文件到 HTML", SHOTS / "ch4" / "02-数据流.png")],
    },
    {
        "anchor": "4.3 安全设计",
        "figures": [(4, 3, "安全分层防御模型", SHOTS / "ch4" / "03-安全分层.png")],
    },
    {
        "anchor": "6.2 线上更新",
        "figures": [(6, 1, "线上内容更新与发布流程", SHOTS / "ch6" / "01-维护流程.png")],
    },
]

HEADING_RE = re.compile(r"^\d+(\.\d+)+\s+\S")
FIG_CAP_RE = re.compile(r"^图(\d+)-(\d+)\s")


def caption_label(ch: int, num: int) -> str:
    return f"图{ch}-{num}"


def is_heading(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    if HEADING_RE.match(t):
        return True
    if t.startswith("第") and "章" in t[:6]:
        return True
    return False


def existing_captions(doc) -> set[str]:
    found: set[str] = set()
    for p in doc.paragraphs:
        m = FIG_CAP_RE.match(p.text.strip())
        if m:
            found.add(f"图{m.group(1)}-{m.group(2)}")
    return found


def find_anchor_index(paragraphs, anchor: str, anchor_alt: str | None = None) -> int | None:
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if anchor in t:
            if anchor_alt is None or anchor_alt in t:
                return i
    if anchor_alt:
        for i, p in enumerate(paragraphs):
            if anchor_alt in p.text.strip():
                return i
    return None


def section_end_index(paragraphs, start: int) -> int:
    """Index of next heading after start, or len(paragraphs)."""
    for i in range(start + 1, len(paragraphs)):
        if is_heading(paragraphs[i].text):
            return i
    return len(paragraphs)


def section_has_caption(paragraphs, start: int, end: int, label: str) -> bool:
    for i in range(start, end):
        if paragraphs[i].text.strip().startswith(label):
            return True
    return False


def insert_paragraph_before(paragraph, doc):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_caption_paragraph(doc, before_para, ch: int, num: int, title: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p = insert_paragraph_before(before_para, doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"图{ch}-{num}  {title}")
    r.font.name = "楷体"
    r.font.size = Pt(10.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    return p


def add_image_paragraph(doc, before_para, image_path: Path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    p = insert_paragraph_before(before_para, doc)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_hint_paragraph(doc, before_para, rel_path: str):
    from docx.shared import Pt

    p = insert_paragraph_before(before_para, doc)
    r = p.add_run(f"（请在此处插入截图：{rel_path}）")
    r.font.size = Pt(10)
    r.italic = True
    return p


def insert_figure_block(doc, before_para, ch: int, num: int, title: str, image_path: Path) -> None:
    rel = image_path.relative_to(HOMEWORK).as_posix() if image_path.is_file() else str(image_path)
    if image_path.is_file():
        add_image_paragraph(doc, before_para, image_path)
    else:
        add_hint_paragraph(doc, before_para, rel)
    add_caption_paragraph(doc, before_para, ch, num, title)


def backup_docx(src: Path) -> Path:
    BACKUP.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest = BACKUP / f"{src.stem}_pre_insert_{ts}.docx"
    shutil.copy2(src, dest)
    return dest


def main() -> None:
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx")
        raise SystemExit(1)

    if not OUT.is_file():
        print("missing", OUT)
        raise SystemExit(1)

    pre = backup_docx(OUT)
    print("backup:", pre)

    doc = Document(str(OUT))
    paragraphs = doc.paragraphs
    captions = existing_captions(doc)

    # collect insertions: (insert_before_index, figures list)
    insertions: list[tuple[int, list[tuple]]] = []
    for slot in FIGURE_SLOTS:
        anchor = slot["anchor"]
        anchor_alt = slot.get("anchor_alt")
        idx = find_anchor_index(paragraphs, anchor, anchor_alt)
        if idx is None:
            print("WARN anchor not found:", anchor)
            continue
        end = section_end_index(paragraphs, idx)
        pending = []
        for ch, num, title, path in slot["figures"]:
            label = caption_label(ch, num)
            if label in captions or section_has_caption(paragraphs, idx, end, label):
                print("skip existing", label)
                continue
            pending.append((ch, num, title, path))
        if pending:
            insertions.append((end, pending))

    # insert from bottom to top so indices stay valid
    insertions.sort(key=lambda x: x[0], reverse=True)
    inserted = 0
    for before_idx, figures in insertions:
        if before_idx >= len(doc.paragraphs):
            before_para = doc.paragraphs[-1]
        else:
            before_para = doc.paragraphs[before_idx]
        for ch, num, title, path in reversed(figures):
            insert_figure_block(doc, before_para, ch, num, title, path)
            inserted += 1
            print("inserted", caption_label(ch, num), "before idx", before_idx)

    try:
        doc.save(str(OUT))
        print(f"saved: {OUT}")
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_插图已更新.docx")
        doc.save(str(alt))
        print(f"原文件被 Word 占用，已另存为: {alt}")
        print("请关闭 Word 后将该文件覆盖原 docx。")
    print(f"inserted {inserted} figure block(s)")


if __name__ == "__main__":
    main()
